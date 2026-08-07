from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


class KnowledgeDocumentIngestionError(Exception):
    """
    Knowledge Base dosya içe aktarma hatasıdır.
    """


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    filename: str
    extension: str


def _normalize_text(text: str) -> str:
    lines = [
        line.strip()
        for line in text.replace("\r\n", "\n").split("\n")
    ]

    normalized_lines = []
    previous_blank = False

    for line in lines:
        if not line:
            if previous_blank:
                continue

            normalized_lines.append("")
            previous_blank = True
            continue

        normalized_lines.append(line)
        previous_blank = False

    return "\n".join(normalized_lines).strip()


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(
        BytesIO(content)
    )

    pages = []

    for page in reader.pages:
        text = page.extract_text() or ""

        if text.strip():
            pages.append(text)

    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    document = DocxDocument(
        BytesIO(content)
    )

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    return "\n\n".join(paragraphs)


def _extract_txt(content: bytes) -> str:
    for encoding in (
        "utf-8-sig",
        "utf-8",
        "cp1254",
        "latin-1",
    ):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise KnowledgeDocumentIngestionError(
        "TXT dosyasının karakter kodlaması çözümlenemedi."
    )


def extract_document_text(
    *,
    filename: str,
    content: bytes,
) -> ExtractedDocument:
    extension = Path(
        filename
    ).suffix.lower()

    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise KnowledgeDocumentIngestionError(
            "Yalnızca PDF, DOCX veya TXT "
            "dokümanları desteklenmektedir."
        )

    if not content:
        raise KnowledgeDocumentIngestionError(
            "Yüklenen doküman boş."
        )

    try:
        if extension == ".pdf":
            text = _extract_pdf(content)

        elif extension == ".docx":
            text = _extract_docx(content)

        else:
            text = _extract_txt(content)

    except KnowledgeDocumentIngestionError:
        raise

    except Exception as error:
        raise KnowledgeDocumentIngestionError(
            "Doküman içeriği okunamadı."
        ) from error

    normalized_text = _normalize_text(text)

    if not normalized_text:
        raise KnowledgeDocumentIngestionError(
            "Dokümandan indekslenebilir metin çıkarılamadı."
        )

    return ExtractedDocument(
        text=normalized_text,
        filename=filename,
        extension=extension,
    )
